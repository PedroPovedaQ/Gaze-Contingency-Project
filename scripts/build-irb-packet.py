#!/usr/bin/env python3
"""Build editable IRB packet DOCX files from auditable Markdown sources."""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "irb" / "source"
DEFAULT_OUTPUT = ROOT / "docs" / "irb" / "submission"
PROTOCOL_TEMPLATE = ROOT / "docs" / "irb" / "templates" / "HRP-503 - TEMPLATE - Protocol.docx"
CONSENT_TEMPLATE = ROOT / "docs" / "irb" / "templates" / "HRP-502 - TEMPLATE CONSENT DOCUMENT - Adult.docx"

DOCUMENTS = [
    ("protocol.md", "HRP-503_Gaze_Contingency_Protocol_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("consent.md", "HRP-502_Gaze_Contingency_Adult_Consent_DRAFT.docx", CONSENT_TEMPLATE),
    ("recruitment.md", "Recruitment_Materials_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("screening.md", "Eligibility_and_Safety_Screening_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("questionnaires.md", "Participant_Questionnaires_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("voice-script.md", "Voice_Recording_Script_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("session-record.md", "Researcher_Session_and_Event_Record_DRAFT.docx", PROTOCOL_TEMPLATE),
    ("post-participation.md", "Post_Participation_Information_DRAFT.docx", PROTOCOL_TEMPLATE),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def preserve_logo_and_clear(document: Document) -> None:
    body = document.element.body
    logo = None
    for child in body.iterchildren():
        if child.tag.endswith("}p") and child.xpath(".//w:drawing"):
            logo = deepcopy(child)
            break

    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    if logo is not None:
        for doc_pr in logo.xpath(".//wp:docPr"):
            doc_pr.set("title", "University of Central Florida logo")
            doc_pr.set("descr", "University of Central Florida logo")
        body.insert(0, logo)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color in (
        ("Title", 18, "000000"),
        ("Heading 1", 14, "1F4E79"),
        ("Heading 2", 11, "1F4E79"),
        ("Heading 3", 10, "404040"),
    ):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    for name in ("List Bullet", "List Number"):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)

    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.start_type = WD_SECTION_START.CONTINUOUS


def add_inline(paragraph, text: str, bold_default: bool = False) -> None:
    """Render **bold** and highlight OPEN DECISION spans."""
    parts = re.split(r"(\*\*.*?\*\*|\[OPEN DECISION[^\]]*\])", text)
    for part in parts:
        if not part:
            continue
        bold = bold_default
        value = part
        if part.startswith("**") and part.endswith("**"):
            value = part[2:-2]
            bold = True
        run = paragraph.add_run(value)
        run.bold = bold
        run.font.name = "Arial"
        if value.startswith(("[OPEN DECISION", "OPEN DECISION")):
            run.bold = True
            run.font.color.rgb = RGBColor(156, 0, 6)
            run.font.highlight_color = 7  # yellow


def add_draft_banner(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FCE4D6")
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "F4B084")
        p_bdr.append(border)
    p_pr.append(p_bdr)
    run = p.add_run("DRAFT — NOT FOR SUBMISSION OR PARTICIPANT USE\nResolve all OPEN DECISION items and obtain UCF approval before use.")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(156, 0, 6)


def is_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        if not is_separator(lines[idx]):
            rows.append([cell.strip() for cell in lines[idx].strip().strip("|").split("|")])
        idx += 1
    return rows, idx


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        tr_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if r_idx == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for c_idx in range(width):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, row[c_idx] if c_idx < len(row) else "", bold_default=r_idx == 0)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown(document: Document, source_text: str) -> None:
    lines = source_text.splitlines()
    idx = 0
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = document.add_paragraph()
            add_inline(p, " ".join(x.strip() for x in paragraph_buffer))
            paragraph_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if not stripped:
            flush()
            idx += 1
            continue
        if stripped == "<!-- PAGE BREAK -->":
            flush()
            document.add_page_break()
            idx += 1
            continue
        if stripped.startswith("|"):
            flush()
            rows, idx = parse_table(lines, idx)
            add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush()
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                p = document.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = document.add_paragraph(style=f"Heading {level - 1}")
            add_inline(p, text, bold_default=True)
            idx += 1
            continue
        bullet = re.match(r"^\s*-\s+(.*)$", raw)
        numbered = re.match(r"^\s*(\d+)\.\s+(.*)$", raw)
        if bullet or numbered:
            flush()
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            prefix = "• " if bullet else f"{numbered.group(1)}. "
            content = bullet.group(1) if bullet else numbered.group(2)
            add_inline(p, prefix + content.removesuffix("\\").rstrip())
            idx += 1
            continue
        if raw.endswith("\\"):
            paragraph_buffer.append(stripped[:-1].rstrip())
            flush()
        else:
            paragraph_buffer.append(stripped)
        idx += 1
    flush()


def add_doc_properties(document: Document, title: str) -> None:
    props = document.core_properties
    props.title = title
    props.subject = "First-pass UCF IRB submission draft"
    props.author = "Gaze Contingency Project research team"
    props.keywords = "UCF, IRB, mixed reality, eye tracking, voice cloning"
    props.comments = "Generated from auditable Markdown; resolve all OPEN DECISION items before submission."


def build(source_name: str, output_name: str, template: Path, output_dir: Path) -> None:
    document = Document(template)
    preserve_logo_and_clear(document)
    configure_styles(document)
    add_draft_banner(document)
    add_markdown(document, (SOURCE / source_name).read_text(encoding="utf-8"))
    add_doc_properties(document, output_name.removesuffix(".docx"))

    # Make the source template/version relationship explicit without adding tracked comments.
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Draft generated from {template.name} structure and styles. Content source: {source_name}.")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    output_dir.mkdir(parents=True, exist_ok=True)
    document.save(output_dir / output_name)
    print(output_dir / output_name)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    for source_name, output_name, template in DOCUMENTS:
        build(source_name, output_name, template, args.output_dir.resolve())


if __name__ == "__main__":
    main()
