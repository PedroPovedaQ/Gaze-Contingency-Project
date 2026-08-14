#!/usr/bin/env python3
"""Build editable first-pass IRB documents for the Gaze Contingency Project."""

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "irb" / "project"
BLUE = "173B57"
TEAL = "1F6F78"
LIGHT_BLUE = "EAF2F6"
LIGHT_AMBER = "FFF4D6"
GRAY = "59636B"
LIGHT_GRAY = "F2F4F5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Aptos", size=10.5, bold=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 12.5, TEAL),
        ("Heading 3", 11, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(short_title)
    set_font(r, size=8.5, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gaze Contingency Project | IRB working draft | July 27, 2026")
    set_font(r, size=8, color=GRAY)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, name="Aptos Display", size=23, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    set_font(r, size=11.5, color=TEAL)
    add_callout(
        doc,
        "FIRST-PASS DRAFT",
        "Not IRB-approved. Replace all bracketed open decisions and transfer final "
        "content into the current official UCF forms before submission.",
        LIGHT_AMBER,
    )


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.75)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, size=9, bold=True, color=BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=BLUE)
    p.add_run(text)


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        if widths:
            set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(text))
        set_font(r, size=9, bold=True, color=WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cell = cells[i]
            if row_index % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            set_cell_margins(cell)
            if widths:
                set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_font(r, size=9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_scale(doc: Document, anchors: str, values: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, text, fill in ((0, values, LIGHT_BLUE), (1, anchors, WHITE)):
        cell = table.cell(row, 0)
        set_cell_width(cell, 6.65)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 80, 100, 80, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=8.5, color=GRAY if row else BLUE, bold=(row == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_question(doc: Document, number: str, text: str, response: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number}. ")
    set_font(r, bold=True, color=TEAL)
    p.add_run(text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(response)
    set_font(r, size=9.5, color=GRAY)


def build_protocol() -> Path:
    doc = Document()
    configure_document(doc, "HRP-503 content draft")
    add_title(
        doc,
        "Gaze-Contingent Guidance in Mixed-Reality Visual Search",
        "Draft content organized for a UCF HRP-503 protocol application",
    )

    doc.add_heading("1. Administrative Summary", level=1)
    add_table(
        doc,
        ["Field", "Draft entry"],
        [
            ("Principal investigator", "[OPEN DECISION: name, department, email, phone]"),
            ("Faculty advisor", "[OPEN DECISION: name, department, email]"),
            ("Study site", "[OPEN DECISION: approved UCF laboratory room]"),
            ("Population", "Adults 18 or older who can safely use an MR headset and complete an English-language visual-search task."),
            ("Local enrollment", "[OPEN DECISION: power/precision analysis and pilot allowance]"),
            ("Session length", "Approximately 45-60 minutes; replace with pilot timing."),
            ("Compensation", "[OPEN DECISION: amount, form, and partial-payment policy]"),
            ("Current executable design", "One within-participant run of 14 rounds alternating gaze-unaware and gaze-aware spoken guidance."),
            ("Future design", "A 2 x 2 gaze-awareness by voice-similarity design is proposed but is not implemented."),
        ],
        widths=[1.65, 5.05],
    )

    doc.add_heading("2. Background and Purpose", level=1)
    doc.add_paragraph(
        "Mixed-reality visual search requires people to distribute attention across "
        "spatially arranged objects while coordinating eye gaze, head pose, target "
        "memory, and selection. The study evaluates whether spoken guidance that can "
        "use gaze-derived interaction state improves search performance or experience "
        "relative to guidance that does not use gaze information."
    )
    doc.add_paragraph(
        "Prior visual-search and gaze research motivates eye calibration, controlled "
        "stimuli, explicit timing boundaries, workload measurement, and careful "
        "preprocessing. The study adapts those practices but does not treat thresholds "
        "from other headsets as validated for the HTC Vive Focus Vision."
    )
    add_labeled_paragraph(
        doc,
        "Primary research question",
        "Does gaze-aware spoken guidance change search completion time relative to "
        "gaze-unaware spoken guidance?",
    )
    add_labeled_paragraph(
        doc,
        "Secondary questions",
        "Do conditions differ in wrong captures, first-try success, subjective "
        "workload, perceived helpfulness, trust, reliance, distraction, intrusiveness, "
        "or gaze-derived exploratory behavior?",
    )

    doc.add_heading("3. Design and Conditions", level=1)
    add_callout(
        doc,
        "CURRENT DESIGN BOUNDARY",
        "The IRB submission must select one executable design. The present build is "
        "the two-condition 14-round study. The voice-similarity factor should remain "
        "out of scope unless it is implemented, separately consented, and fully "
        "specified before submission.",
    )
    add_table(
        doc,
        ["Condition", "Information available", "Participant experience"],
        [
            (
                "Gaze-unaware",
                "Task state and elapsed time; no current gaze/hover evidence.",
                "Generic encouragement delivered by the same audio path.",
            ),
            (
                "Gaze-aware",
                "Task state plus current/recent gaze-hover, shelf, scanning, and target-proximity evidence.",
                "Spoken feedback may reflect whether search appears off-target, on-track, or near the target.",
            ),
        ],
        widths=[1.25, 2.6, 2.85],
    )
    doc.add_paragraph(
        "The causal comparison requires matched hint opportunities, timing, audio "
        "quality, voice identity, and semantic usefulness. The current implementation "
        "uses different default delays/intervals and therefore requires correction or "
        "a narrower pilot interpretation."
    )

    doc.add_heading("4. Participants", level=1)
    doc.add_heading("Inclusion criteria", level=2)
    add_bullets(
        doc,
        [
            "Age 18 or older.",
            "Able to understand the approved English consent and instructions.",
            "Able to wear the HTC Vive Focus Vision and view the search display with normal or corrected-to-normal vision.",
            "Able to distinguish the task colors or pass the selected color-vision screen.",
            "Able to hear spoken guidance at a comfortable level.",
            "Willing to permit coded gaze, interaction, task-performance, and questionnaire data collection.",
        ],
    )
    doc.add_heading("Exclusion or stopping criteria", level=2)
    add_bullets(
        doc,
        [
            "Current dizziness, nausea, severe headache, unusual fatigue, or condition making headset use unsafe.",
            "History of severe simulator sickness, relevant vestibular disorder, or visually triggered seizure risk.",
            "Vision, color discrimination, or hearing limitation that prevents task completion after reasonable correction.",
            "Eye-tracking calibration that remains unusable after refitting and the approved retry procedure.",
            "Participant request to stop, significant discomfort, unsafe tracking, or unrecoverable hardware/software failure.",
        ],
    )
    doc.add_paragraph(
        "Glasses, contact lenses, eye makeup, prior XR experience, and prior eye-tracking "
        "experience should be recorded as descriptors unless piloting demonstrates a "
        "specific safety or data-quality reason for exclusion."
    )

    doc.add_heading("5. Recruitment and Consent", level=1)
    doc.add_paragraph(
        "Recruitment may use approved email, flyer, participant-pool, or word-of-mouth "
        "methods. Materials will describe mixed-reality headset use, eye tracking, "
        "spoken guidance, questionnaires, session duration, eligibility, compensation, "
        "and voluntary participation without describing one condition as expected to be superior."
    )
    doc.add_paragraph(
        "A trained researcher will conduct consent before research data collection, "
        "review procedures and risks, explain gaze-data collection and confidentiality, "
        "answer questions, and confirm that participation may be paused or stopped "
        "without penalty. Written versus verbal consent and any waiver of documentation "
        "remain open decisions for the PI and IRB."
    )

    doc.add_heading("6. Study Procedure", level=1)
    add_numbered(
        doc,
        [
            "Confirm minimal eligibility and review consent.",
            "Assign a coded participant ID; keep scheduling and compensation records separate.",
            "Collect approved baseline demographics, XR experience, and baseline simulator-sickness ratings.",
            "Fit the Vive Focus Vision and complete vendor eye-tracking calibration.",
            "Validate gaze across the intended search area; refit/recalibrate according to the approved rule.",
            "Complete practice using targets and layouts excluded from analyzed trials.",
            "Start the MR task and present the target during a fixation-cross transition.",
            "For each round, display 56 objects, allow gaze-dwell selection, provide condition-dependent spoken guidance, and continue after wrong captures until the correct target is selected.",
            "Provide breaks at the approved intervals and whenever requested.",
            "Collect raw NASA-TLX and the guidance manipulation check after each condition block if a block design is implemented.",
            "After the final task, collect post-SSQ, overall SUS, comparative guidance/privacy items, and a short interview.",
            "Remove the headset, check participant comfort, debrief, confirm compensation, and document deviations.",
        ],
    )

    doc.add_heading("7. Apparatus and Data", level=1)
    add_table(
        doc,
        ["Source", "Recorded data", "Current status"],
        [
            (
                "gaze_log.csv",
                "Rendered-frame timestamp, gaze origin/direction/rotation, available per-eye position/openness/fixation point, hovered object, target match, dwell progress, task state, blink estimate.",
                "Implemented; not a confirmed native-rate stream and lacks complete validity/calibration metadata.",
            ),
            (
                "trial_events.csv",
                "Game/objective events, correct/wrong captures, hover-derived episode boundaries, object metadata, dwell duration.",
                "Implemented; existing fixation/saccade names are interaction proxies.",
            ),
            (
                "trial_summary.json",
                "Per-round timing, wrong captures, first-try success, target/distractor hover summaries, blink summaries.",
                "Implemented; some timing boundaries require correction.",
            ),
            (
                "nasa_tlx.csv",
                "Six 0-100 raw NASA-TLX subscales and aggregate raw score.",
                "Implemented once after the run; condition-specific collection requires runtime changes.",
            ),
            (
                "Study questionnaires",
                "Demographics, XR experience, SSQ, SUS, guidance manipulation checks, privacy/trust, interview.",
                "Protocol proposal; instrument set requires PI/IRB approval.",
            ),
        ],
        widths=[1.25, 3.65, 1.8],
    )

    doc.add_heading("8. Measures and Analysis", level=1)
    add_table(
        doc,
        ["Construct", "Operational definition", "Role and limitation"],
        [
            (
                "Search time",
                "Seconds from dedicated search_onset/display-ready event to correct dwell capture.",
                "Proposed primary outcome; current start boundary must be fixed.",
            ),
            (
                "Selection accuracy",
                "First-attempt correctness and number of wrong captures per round.",
                "Confirmatory or secondary performance outcome.",
            ),
            (
                "Subjective workload",
                "Six raw NASA-TLX subscales, each 0-100, after each condition block.",
                "Primary workload measure; cannot compare conditions from one end-of-run rating.",
            ),
            (
                "Guidance manipulation",
                "Post-block ratings of responsiveness to looking/search behavior and perceived timing.",
                "Manipulation check.",
            ),
            (
                "Hover allocation",
                "Target/distractor hover duration, hover episodes, revisits, and target-entry latency.",
                "Exploratory interaction proxy; not a validated fixation measure.",
            ),
            (
                "Eye behavior",
                "Offline gaze angular velocity, dispersion, and pilot-validated fixation/saccade features.",
                "Exploratory until sampling, quality, and classifier validation are complete.",
            ),
            (
                "Usability and safety",
                "SUS total score and pre/post SSQ symptom scores.",
                "Secondary system and safety outcomes.",
            ),
        ],
        widths=[1.25, 3.55, 1.9],
    )
    doc.add_paragraph(
        "Repeated observations should be analyzed with models that account for "
        "participant, round/order, target identity, and condition. The final plan must "
        "predeclare the primary outcome, multiplicity strategy, missing-data rules, "
        "gaze-quality exclusions, and any participant/round exclusions."
    )

    doc.add_heading("9. Risks and Mitigation", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ("Eye strain, headache, fatigue, headset discomfort", "Fit headset carefully; provide breaks; stop on request or significant symptoms."),
            ("Simulator/MR sickness or disorientation", "Stable passthrough scene, no artificial locomotion, pre/post SSQ, clear physical area, seated option if approved."),
            ("Frustration or embarrassment about performance", "Neutral instructions and feedback; explain that wrong selections are expected; permit skipping optional questions."),
            ("Gaze-data privacy", "Coded IDs, data minimization, restricted approved storage, separate linkage data, aggregate reporting."),
            ("Audio/network interruption", "Preflight test, matched fallback behavior, deviation logging, no pressure to repeat after discomfort."),
            ("Future voice cloning", "Out of current scope unless separately consented with provider, retention, deletion, and reuse limits."),
        ],
        widths=[2.25, 4.45],
    )

    doc.add_heading("10. Data Management", level=1)
    doc.add_paragraph(
        "Research files will use coded participant IDs. Names, contact details, consent "
        "records, scheduling information, and compensation records will be stored "
        "separately from task data. Files will be transferred from the headset/study "
        "computer to [OPEN DECISION: approved UCF-managed storage], verified, and local "
        "working copies removed according to the approved procedure. Access, encryption, "
        "retention duration, deletion, withdrawal cutoff, and de-identified sharing "
        "language remain open decisions."
    )
    doc.add_paragraph(
        "Gaze telemetry is sensitive behavioral data. It will not be described as "
        "diagnostic, and reports will use aggregate statistics or de-identified examples."
    )

    doc.add_heading("11. Required Pre-Submission Resolutions", level=1)
    add_bullets(
        doc,
        [
            "Select the executable study design and remove unsupported voice-factor language if out of scope.",
            "Match hint opportunities/timing/content and add a complete hint audit trail.",
            "Add a valid search_onset event.",
            "Counterbalance condition and target/layout order.",
            "Select sample size using an a priori power/precision analysis.",
            "Finalize compensation, site, study duration, and stopping rules.",
            "Finalize calibration validation, achieved-sampling, validity, missingness, and exclusion rules.",
            "Implement condition-specific NASA-TLX or limit workload claims.",
            "Finalize data storage, retention, deletion, access, sharing, and withdrawal language.",
            "Have the PI/faculty advisor verify instrument permissions, scoring, and current UCF template requirements.",
        ],
    )

    path = OUT / "gaze-contingency-hrp-503-content-draft.docx"
    doc.save(path)
    return path


def build_participant_materials() -> Path:
    doc = Document()
    configure_document(doc, "Participant materials")
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.03
    add_title(
        doc,
        "Participant-Facing Materials",
        "Recruitment, screening, consent-information, session script, and debrief drafts",
    )

    doc.add_heading("A. Recruitment Email", level=1)
    add_labeled_paragraph(doc, "Subject", "Mixed-Reality Visual Search Study Recruitment")
    doc.add_paragraph(
        "Researchers at the University of Central Florida are inviting adults to "
        "participate in a study about mixed-reality object search and computer guidance."
    )
    doc.add_paragraph(
        "Participation involves one in-person session lasting approximately "
        "[OPEN DECISION: duration]. You will wear an HTC Vive Focus Vision headset with "
        "eye tracking, search for virtual objects on a shelf-like display, hear spoken "
        "guidance, and answer questionnaires. You may experience temporary eye strain, "
        "headset discomfort, fatigue, dizziness, nausea, or frustration."
    )
    add_bullets(
        doc,
        [
            "You must be at least 18 years old.",
            "You must be able to understand the approved English instructions and consent information.",
            "You must be able to safely wear a mixed-reality headset, see the task colors and objects with normal or corrected vision, and hear spoken guidance.",
            "You should not participate if you currently feel dizzy, nauseous, unwell, or unusually fatigued, or if you have a condition that makes headset use unsafe.",
        ],
    )
    doc.add_paragraph(
        "Compensation: [OPEN DECISION]. Participation is voluntary, and choosing not "
        "to participate will not affect your grades, employment, services, or "
        "relationship with UCF."
    )
    doc.add_paragraph("Study contact: [OPEN DECISION: approved PI name and contact]")
    doc.add_paragraph("Scheduling link or QR code: [OPEN DECISION]")

    doc.add_heading("B. Minimal Eligibility Prescreen", level=1)
    add_callout(
        doc,
        "ADMINISTRATION NOTE",
        "Collect only what is needed to determine eligibility before consent. Do not "
        "retain medical detail or research questionnaire responses from people who do "
        "not consent unless the approved protocol explicitly permits it.",
    )
    questions = [
        ("1", "Are you 18 years of age or older?", "Yes / No"),
        ("2", "Can you read and understand the approved English consent and task instructions?", "Yes / No"),
        ("3", "Can you safely wear a mixed-reality headset for approximately [duration] with breaks?", "Yes / No / Unsure"),
        ("4", "Do you currently feel dizzy, nauseous, unwell, have a severe headache, or feel unusually fatigued?", "Yes / No / Prefer not to say"),
        ("5", "Do you have a history of severe simulator sickness, a relevant vestibular disorder, or visually triggered seizures that would make headset use unsafe?", "Yes / No / Prefer not to say"),
        ("6", "With your usual correction, can you see and distinguish the task objects/colors and hear spoken guidance?", "Yes / No / Unsure"),
        ("7", "Are you willing to allow coded gaze, interaction, task-performance, and questionnaire data collection?", "Yes / No"),
    ]
    for number, text, response in questions:
        add_question(doc, number, text, response)
    doc.add_paragraph(
        "Eligibility outcome: Eligible / Not eligible / PI review required. Record only "
        "the outcome and approved minimal screening fields."
    )

    doc.add_heading("C. Participant Information and Consent Draft", level=1)
    add_callout(
        doc,
        "TEMPLATE NOTE",
        "This is content for review, not a substitute for the current official UCF "
        "consent template. The PI and IRB must determine written versus verbal consent.",
    )
    add_labeled_paragraph(doc, "Study title", "Gaze-Contingent Guidance in Mixed-Reality Visual Search")
    add_labeled_paragraph(doc, "Research team", "[OPEN DECISION: PI, faculty advisor, department, contacts]")

    doc.add_heading("Why am I being invited?", level=2)
    doc.add_paragraph(
        "You are being invited because you are an adult who may be able to safely use "
        "a mixed-reality headset and complete a visual-search task."
    )
    doc.add_heading("Why is this research being done?", level=2)
    doc.add_paragraph(
        "The purpose is to understand whether spoken computer guidance that uses "
        "information about where a person is looking changes object-search performance, "
        "workload, and experience compared with guidance that does not use gaze information."
    )
    doc.add_heading("What will happen?", level=2)
    doc.add_paragraph(
        "If you agree, you will complete brief background and safety questions, wear "
        "an HTC Vive Focus Vision mixed-reality headset, complete eye-tracking "
        "calibration and practice, and search for announced target objects among "
        "virtual distractors. You select an object by looking at it continuously until "
        "the dwell indicator completes. The system may provide spoken guidance. Some "
        "guidance may use current or recent gaze behavior, while other guidance does "
        "not. You may complete workload, sickness, usability, and experience questions "
        "and a short interview. The visit will last approximately [OPEN DECISION]."
    )
    doc.add_heading("What information will be collected?", level=2)
    doc.add_paragraph(
        "The study may collect a coded participant ID, task condition and configuration, "
        "gaze origin and direction, available eye openness and device fixation-point "
        "signals, object hover and dwell behavior, target and selected object metadata, "
        "timing, wrong and correct selections, estimated blinks when available, spoken-"
        "guidance events, calibration/quality information, questionnaire responses, and "
        "researcher-recorded technical deviations. The current study does not intend to "
        "record your face, room video, or ordinary conversation."
    )
    doc.add_heading("What are the risks?", level=2)
    doc.add_paragraph(
        "Possible temporary risks include eye strain, headache, fatigue, warmth or "
        "pressure from the headset, dizziness, nausea, disorientation, frustration, "
        "and discomfort about eye-tracking privacy. You may pause, request a break, "
        "skip optional questions, or stop at any time. Tell the researcher immediately "
        "if you feel uncomfortable."
    )
    doc.add_heading("Are there benefits?", level=2)
    doc.add_paragraph(
        "You may not receive a direct personal benefit. The research may contribute to "
        "the design of mixed-reality systems that provide better-timed or more relevant assistance."
    )
    doc.add_heading("How will privacy be protected?", level=2)
    doc.add_paragraph(
        "Research files will use a coded participant ID. Identifying scheduling, "
        "consent, and compensation information will be stored separately from research "
        "data. Access will be limited to approved study personnel. Data will be stored "
        "on [OPEN DECISION: approved UCF-managed storage] for [OPEN DECISION: retention "
        "period] and then handled according to the approved deletion or archiving plan. "
        "Complete secrecy cannot be guaranteed; authorized UCF or IRB representatives "
        "may inspect study records."
    )
    doc.add_heading("Do I have to participate?", level=2)
    doc.add_paragraph(
        "No. Participation is voluntary. You may decline, stop, or withdraw without "
        "penalty or loss of benefits to which you are otherwise entitled. The approved "
        "form will explain whether and until when already-collected data can be removed."
    )
    doc.add_heading("Compensation and contacts", level=2)
    doc.add_paragraph(
        "Compensation: [OPEN DECISION: amount, form, completion and partial-payment "
        "policy]. Study questions: [PI/faculty advisor contacts]. Questions about your "
        "rights as a research participant: [current UCF IRB contact language from the "
        "official template]."
    )

    doc.add_heading("D. Standardized Session Script", level=1)
    add_numbered(
        doc,
        [
            "Welcome the participant, confirm eligibility, review consent, answer questions, and obtain approved consent.",
            "Assign the coded ID and begin the session manifest.",
            "Administer baseline questionnaires and baseline SSQ.",
            "Explain that breaks and stopping are always available and that wrong selections are expected.",
            "Fit the headset, adjust for comfort, and complete vendor eye calibration.",
            "Validate gaze across the search area. Refit and recalibrate if the approved criterion is not met.",
            "Explain: find the announced color-shape target as quickly and accurately as possible; select by continuous gaze dwell; continue after a wrong selection; spoken hints may be used or ignored.",
            "Run practice until the approved criterion is met without experimenter intervention.",
            "Run the assigned condition schedule. Do not reveal condition labels or hypotheses.",
            "At approved breaks, check comfort using neutral language and administer scheduled post-block measures.",
            "After the task, administer post-SSQ, SUS, comparative questions, and interview before revealing the manipulation.",
            "Remove the headset, check comfort, debrief, confirm compensation, and record deviations or withdrawal requests.",
        ],
    )

    doc.add_heading("E. Stopping and Adverse-Event Script", level=1)
    doc.add_paragraph(
        "Stop the task immediately if the participant asks to stop or reports significant "
        "nausea, dizziness, disorientation, headache, eye strain, distress, or other "
        "concerning symptoms. Pause the application, assist with safe headset removal, "
        "offer a seated rest, and follow the approved escalation/reporting procedure. "
        "Do not encourage the participant to continue. Record the coded event and only "
        "the minimum approved detail."
    )

    doc.add_heading("F. Debriefing Script", level=1)
    debrief = doc.add_paragraph(
        "Thank you for participating. This study compares spoken guidance that can use "
        "gaze information with guidance that does not use gaze information. In gaze-aware "
        "conditions, the assistant could use where and how you were looking to choose "
        "feedback. In gaze-unaware conditions, it used general guidance without gaze "
        "evidence. We are interested in whether this changes search performance, "
        "workload, trust, helpfulness, or privacy perceptions. Your coded data will be "
        "handled as described in the consent information. Please contact the research "
        "team if you have questions or wish to request withdrawal within the approved timeframe."
    )
    debrief.paragraph_format.keep_together = True

    path = OUT / "gaze-contingency-participant-materials-draft.docx"
    doc.save(path)
    return path


def build_questionnaires() -> Path:
    doc = Document()
    configure_document(doc, "Questionnaire packet")
    add_title(
        doc,
        "Questionnaire Packet",
        "Proposed screening, baseline, post-block, and post-session measures",
    )
    add_callout(
        doc,
        "MEASUREMENT BOUNDARY",
        "Validated instruments are identified by name and should retain approved "
        "wording/scoring. Guidance, trust, reliance, and privacy items are "
        "study-specific and must not be represented as validated scales.",
    )

    doc.add_heading("Administration Schedule", level=1)
    add_table(
        doc,
        ["Time", "Measures", "Reason"],
        [
            ("Eligibility", "Minimal prescreen", "Safety and core eligibility only"),
            ("After consent", "Demographics, XR/eye-tracking experience, baseline SSQ", "Sample description, covariates, symptom baseline"),
            ("After each condition block", "Raw NASA-TLX and guidance manipulation/experience items", "Condition-level workload and manipulation check"),
            ("After final block", "Post-SSQ, SUS, comparative/privacy items, interview", "Safety change, overall usability, interpretation"),
        ],
        widths=[1.3, 3.15, 2.25],
    )

    doc.add_heading("1. Participant and Session Header", level=1)
    add_question(doc, "1.1", "Coded participant ID", "________________")
    add_question(doc, "1.2", "Session date/time", "________________")
    add_question(doc, "1.3", "Application build and condition schedule ID", "________________")
    add_question(doc, "1.4", "Questionnaire time point", "Baseline / Block ___ / Post-session")

    doc.add_heading("2. Demographics and Relevant Background", level=1)
    add_callout(
        doc,
        "OPTIONAL RESPONSES",
        "Except for eligibility fields approved as required, demographic questions "
        "should permit Prefer not to answer and should be collected only when they "
        "serve sample description, safety, or prespecified analysis.",
    )
    demographic_questions = [
        ("2.1", "Age", "18-24 / 25-34 / 35-44 / 45-54 / 55-64 / 65+ / Prefer not to answer"),
        ("2.2", "Gender", "Woman / Man / Non-binary / Self-describe: ______ / Prefer not to answer"),
        ("2.3", "Current academic or employment status", "Undergraduate / Graduate / Faculty or staff / Not currently affiliated with UCF / Other / Prefer not to answer"),
        ("2.4", "Dominant hand", "Right / Left / Ambidextrous / Prefer not to answer"),
        ("2.5", "Vision used during the study", "Uncorrected / Glasses / Contact lenses / Other / Prefer not to answer"),
        ("2.6", "Known color-vision deficiency", "No / Yes / Unsure / Prefer not to answer"),
        ("2.7", "Hearing correction used during the study", "None / Hearing aid or other correction / Other / Prefer not to answer"),
        ("2.8", "Prior eye-tracking experience", "None / 1-2 sessions / 3-5 sessions / More than 5 sessions"),
    ]
    for number, text, response in demographic_questions:
        add_question(doc, number, text, response)

    doc.add_heading("3. Prior Experience", level=1)
    add_scale(doc, "None                                               Extensive", "1        2        3        4        5")
    for idx, item in enumerate(
        [
            "Using virtual-reality or mixed-reality headsets",
            "Using gaze or eye tracking as an input method",
            "Playing video games",
            "Completing visual-search or object-finding tasks",
            "Interacting with voice assistants",
        ],
        start=1,
    ):
        add_question(doc, f"3.{idx}", item, "1 / 2 / 3 / 4 / 5")

    doc.add_heading("4. Simulator Sickness Questionnaire (SSQ)", level=1)
    doc.add_paragraph(
        "Rate how much you experience each symptom right now. Administer before headset "
        "exposure and again immediately after the final headset exposure."
    )
    add_scale(doc, "0 = None     1 = Slight     2 = Moderate     3 = Severe", "0        1        2        3")
    ssq_items = [
        "General discomfort",
        "Fatigue",
        "Headache",
        "Eye strain",
        "Difficulty focusing",
        "Increased salivation",
        "Sweating",
        "Nausea",
        "Difficulty concentrating",
        "Fullness of head",
        "Blurred vision",
        "Dizziness with eyes open",
        "Dizziness with eyes closed",
        "Vertigo",
        "Stomach awareness",
        "Burping",
    ]
    add_table(
        doc,
        ["Item", "Baseline 0-3", "Post 0-3"],
        [(item, "", "") for item in ssq_items],
        widths=[4.55, 1.05, 1.05],
    )
    doc.add_paragraph(
        "Scoring and stopping thresholds: [OPEN DECISION: use the approved SSQ scoring "
        "procedure and prespecified safety rule; do not invent a cutoff after data collection.]"
    )

    doc.add_page_break()
    doc.add_heading("5. Raw NASA-TLX", level=1)
    doc.add_paragraph(
        "Complete immediately after each condition block. Rate the block you just "
        "completed, not the session overall. Each scale ranges from 0 to 100 in steps of 5."
    )
    nasa_rows = [
        ("Mental Demand", "How mentally demanding was the block?", "Very low", "Very high"),
        ("Physical Demand", "How physically demanding was the block?", "Very low", "Very high"),
        ("Temporal Demand", "How hurried or rushed did you feel?", "Very low", "Very high"),
        ("Performance", "How unsuccessful did you feel in accomplishing the task?", "Perfect", "Failure"),
        ("Effort", "How hard did you have to work to accomplish your level of performance?", "Very low", "Very high"),
        ("Frustration", "How insecure, discouraged, irritated, stressed, or annoyed were you?", "Very low", "Very high"),
    ]
    add_table(
        doc,
        ["Subscale", "Prompt", "Low anchor", "High anchor", "0-100"],
        [(*row, "") for row in nasa_rows],
        widths=[1.25, 3.0, 0.85, 0.85, 0.7],
    )
    doc.add_paragraph(
        "Raw TLX score: mean of the six subscale ratings (or sum if the analysis "
        "pipeline explicitly retains its current 0-600 convention). Store all six "
        "subscales; do not retain only the aggregate."
    )

    doc.add_heading("6. Post-Block Guidance Measures", level=1)
    doc.add_paragraph(
        "Study-specific items. Rate only the block just completed. These are "
        "manipulation checks and experience items, not a validated scale."
    )
    add_scale(doc, "1 = Strongly disagree                         7 = Strongly agree", "1     2     3     4     5     6     7")
    guidance_items = [
        ("6.1", "The spoken guidance seemed responsive to where or how I was looking."),
        ("6.2", "The spoken guidance arrived at useful moments."),
        ("6.3", "The spoken guidance helped me search more effectively."),
        ("6.4", "I trusted the spoken guidance."),
        ("6.5", "I relied on the spoken guidance when deciding where to search."),
        ("6.6", "The spoken guidance distracted me from the search task."),
        ("6.7", "The spoken guidance felt intrusive."),
        ("6.8", "I felt in control of whether to follow or ignore the guidance."),
        ("6.9", "I understood what the spoken guidance meant."),
    ]
    for number, text in guidance_items:
        add_question(doc, number, text, "1 / 2 / 3 / 4 / 5 / 6 / 7")
    add_question(
        doc,
        "6.10",
        "Without being told the condition name, what information do you think the assistant used to generate its guidance?",
        "________________________________________________________________",
    )

    doc.add_heading("7. System Usability Scale (SUS)", level=1)
    doc.add_paragraph(
        "Administer once after the complete system experience. Use standard SUS scoring "
        "with alternating positive and negative items."
    )
    add_scale(doc, "1 = Strongly disagree                    5 = Strongly agree", "1        2        3        4        5")
    sus_items = [
        "I think that I would like to use this system frequently.",
        "I found the system unnecessarily complex.",
        "I thought the system was easy to use.",
        "I think that I would need the support of a technical person to be able to use this system.",
        "I found the various functions in this system were well integrated.",
        "I thought there was too much inconsistency in this system.",
        "I would imagine that most people would learn to use this system very quickly.",
        "I found the system very cumbersome to use.",
        "I felt very confident using the system.",
        "I needed to learn a lot of things before I could get going with this system.",
    ]
    for idx, item in enumerate(sus_items, start=1):
        add_question(doc, f"7.{idx}", item, "1 / 2 / 3 / 4 / 5")

    doc.add_heading("8. Post-Session Comparison and Privacy", level=1)
    add_scale(doc, "1 = Strongly disagree                         7 = Strongly agree", "1     2     3     4     5     6     7")
    post_items = [
        ("8.1", "I was comfortable with the system using eye-tracking data during the task."),
        ("8.2", "I understood what kinds of gaze information the system was collecting."),
        ("8.3", "I would be comfortable using a gaze-aware assistant in a similar application."),
        ("8.4", "The potential benefit of gaze-aware guidance outweighed my privacy concerns."),
        ("8.5", "I noticed meaningful differences in the guidance across parts of the study."),
    ]
    for number, text in post_items:
        add_question(doc, number, text, "1 / 2 / 3 / 4 / 5 / 6 / 7")
    add_question(
        doc,
        "8.6",
        "Which part of the study had the most useful guidance, if any? Describe it without using condition labels.",
        "________________________________________________________________",
    )

    doc.add_heading("9. Semi-Structured Interview", level=1)
    interview_items = [
        "What strategy did you use to find the announced targets?",
        "When did you choose to follow or ignore the spoken guidance?",
        "Which guidance, if any, felt most helpful? What made it helpful?",
        "Did any guidance feel distracting, poorly timed, inaccurate, or intrusive?",
        "Did you notice the assistant responding to where or how you looked? Please describe.",
        "How did selecting objects by sustained gaze affect your search strategy?",
        "Did eye tracking or gaze-data collection create any privacy or comfort concerns?",
        "What part of the task felt easiest?",
        "What part felt most difficult or confusing?",
        "Is there anything else you would change about the task or assistant?",
    ]
    for idx, item in enumerate(interview_items, start=1):
        add_question(doc, f"9.{idx}", item, "Notes: __________________________________________________________")

    doc.add_heading("10. Optional Future Voice-Similarity Appendix", level=1)
    add_callout(
        doc,
        "OUT OF CURRENT SCOPE",
        "Use only if self-similar voice is implemented and separately approved. Voice "
        "recording/synthesis requires explicit consent, provider disclosure, retention "
        "and deletion rules, and a matched generic-voice control.",
        LIGHT_AMBER,
    )
    voice_items = [
        ("10.1", "The assistant's voice sounded similar to my own voice."),
        ("10.2", "The assistant's voice felt familiar."),
        ("10.3", "The assistant's voice was easy to understand."),
        ("10.4", "The assistant's voice made the guidance more trustworthy."),
        ("10.5", "The assistant's voice made the gaze-aware guidance feel more intrusive."),
    ]
    for number, text in voice_items:
        add_question(doc, number, text, "1 / 2 / 3 / 4 / 5 / 6 / 7")

    path = OUT / "gaze-contingency-questionnaires-draft.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [build_protocol(), build_participant_materials(), build_questionnaires()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
