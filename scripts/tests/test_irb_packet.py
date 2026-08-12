#!/usr/bin/env python3
"""Focused regression tests for the IRB packet generator and verifier."""

from __future__ import annotations

import importlib.util
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[2]


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


builder = load_module("irb_builder", ROOT / "scripts" / "build-irb-packet.py")
verifier = load_module("irb_verifier", ROOT / "scripts" / "verify-irb-packet.py")


class GeneratorTests(unittest.TestCase):
    def test_open_decision_is_highlighted_without_flattening_heading_size(self) -> None:
        document = Document()
        builder.configure_styles(document)
        paragraph = document.add_paragraph(style="Heading 1")
        builder.add_inline(paragraph, "Heading **[OPEN DECISION: resolve]**")

        heading_run, decision_run = paragraph.runs
        self.assertIsNone(heading_run.font.size)
        self.assertIsNone(decision_run.font.size)
        self.assertTrue(decision_run.bold)
        self.assertEqual(str(decision_run.font.color.rgb), "9C0006")
        self.assertEqual(int(decision_run.font.highlight_color), 7)

    def test_trailing_backslash_is_removed_from_numbered_and_bulleted_items(self) -> None:
        document = Document()
        builder.configure_styles(document)
        builder.add_markdown(document, "1. First item \\\n- Second item \\")
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertNotIn("\\", text)


class VerifierTests(unittest.TestCase):
    def test_revision_detection_is_namespace_prefix_independent(self) -> None:
        namespace = b"http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for tag in (
            b"ins", b"del", b"rPrChange", b"pPrChange", b"moveFrom", b"moveTo",
            b"tblGridChange", b"cellIns", b"moveFromRangeStart", b"customXmlDelRangeEnd",
        ):
            xml = b"<x:document xmlns:x=\"" + namespace + b"\"><x:" + tag + b"/></x:document>"
            self.assertTrue(verifier.contains_revision(xml))
        w14 = b"http://schemas.microsoft.com/office/word/2010/wordml"
        self.assertTrue(verifier.contains_revision(b"<x:conflictIns xmlns:x=\"" + w14 + b"\"/>"))
        self.assertFalse(verifier.contains_revision(b"<x:ins xmlns:x=\"https://example.test/not-word\"/>"))
        self.assertTrue(verifier.contains_revision(b"<malformed"))

    def test_open_decisions_compare_exact_text_and_multiplicity(self) -> None:
        source = "[OPEN DECISION: sample size] [OPEN DECISION: provider]"
        output = "[OPEN DECISION: sample size] [OPEN DECISION: sample size]"
        self.assertNotEqual(verifier.open_decisions(source), verifier.open_decisions(output))

    def test_header_residue_is_included_in_document_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "header.docx"
            document = Document()
            document.sections[0].header.paragraphs[0].text = "Meta Quest"
            document.save(path)
            self.assertIn("Meta Quest", verifier.document_text(path))

    def test_verify_package_rejects_formatting_revision(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original = Path(temp_dir) / "draft.docx"
            mutated = Path(temp_dir) / "mutated.docx"
            document = Document()
            document.add_paragraph("DRAFT — NOT FOR SUBMISSION OR PARTICIPANT USE")
            document.save(original)

            with zipfile.ZipFile(original) as source, zipfile.ZipFile(mutated, "w") as target:
                for info in source.infolist():
                    data = source.read(info.filename)
                    if info.filename == "word/document.xml":
                        data = data.replace(b"</w:body>", b"<x:rPrChange xmlns:x=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" x:id=\"1\"/></w:body>")
                    target.writestr(info, data)

            verifier.SOURCE_FOR_OUTPUT[mutated.name] = "session-record.md"
            failures = verifier.verify_package(mutated)
            self.assertTrue(any("tracked changes" in failure for failure in failures))


if __name__ == "__main__":
    unittest.main()
