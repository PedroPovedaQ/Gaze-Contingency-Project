#!/usr/bin/env python3
"""Verify that generated IRB drafts are complete, editable, and review-gated."""

from __future__ import annotations

import argparse
from collections import Counter
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "docs" / "irb" / "submission"
SOURCE = ROOT / "docs" / "irb" / "source"

EXPECTED = {
    "HRP-503_Gaze_Contingency_Protocol_DRAFT.docx",
    "HRP-502_Gaze_Contingency_Adult_Consent_DRAFT.docx",
    "Recruitment_Materials_DRAFT.docx",
    "Eligibility_and_Safety_Screening_DRAFT.docx",
    "Participant_Questionnaires_DRAFT.docx",
    "Voice_Recording_Script_DRAFT.docx",
    "Researcher_Session_and_Event_Record_DRAFT.docx",
    "Post_Participation_Information_DRAFT.docx",
}

SOURCE_FOR_OUTPUT = {
    "HRP-503_Gaze_Contingency_Protocol_DRAFT.docx": "protocol.md",
    "HRP-502_Gaze_Contingency_Adult_Consent_DRAFT.docx": "consent.md",
    "Recruitment_Materials_DRAFT.docx": "recruitment.md",
    "Eligibility_and_Safety_Screening_DRAFT.docx": "screening.md",
    "Participant_Questionnaires_DRAFT.docx": "questionnaires.md",
    "Voice_Recording_Script_DRAFT.docx": "voice-script.md",
    "Researcher_Session_and_Event_Record_DRAFT.docx": "session-record.md",
    "Post_Participation_Information_DRAFT.docx": "post-participation.md",
}

PROTOCOL_HEADINGS = {
    "Version History",
    "Study Summary",
    "Background",
    "Objectives",
    "What is Being Measured or Evaluated?",
    "Drugs and Medical Devices (If not a drug or medical device study, check N/A and skip to the Endpoints section)",
    "Endpoints (clinical trials only—if not a clinical trial, check N/A and skip to the Number of Subjects section)",
    "Number of Subjects",
    "Inclusion/Exclusion Criteria",
    "Recruitment Methods",
    "Consent Process",
    "HIPAA (From medical records only. If data is not coming from medical records, check N/A and skip to the FERPA section)",
    "FERPA (From educational records only. If data is not coming from educational records, check N/A and skip to the Study Procedures section)",
    "Study Procedures (after the consent process)",
    "Setting",
    "Study Timelines",
    "Resources Available",
    "Potential Benefits to Subjects",
    "Risks to Subjects",
    "Compensation for Research-Related Injury (for more than minimal risk studies only—if this is a minimal risk study, check N/A and skip to the Withdrawal of Subjects section)",
    "Withdrawal of Subjects",
    "Economic Burden to Subjects (If there are no economic burdens to subjects, check N/A and skip to the Compensation section)",
    "Compensation",
    "Data Validity",
    "Sharing Results with Subjects",
    "Data Management and Confidentiality",
    "Long-Term Data and/or Specimen Banking (If there are no plans for long-term data and/or specimen banking, check N/A and skip to the Data Analysis Plan section)",
    "Data Analysis Plan",
    "Provisions to Monitor the Data to Ensure the Safety of Subjects (for more than minimal risk studies only—if this is a minimal risk study, check N/A and skip to the Provisions to Protect the Privacy Interests of Subjects section)",
    "Provisions to Protect the Privacy Interests of Subjects",
    "Adverse Events, Unanticipated Problems, and Deviations",
    "Records",
    "Rationale for the Voice-Cloning Factor and Provider Path",
    "Multi-Site Research",
    "References and External Services",
}

CONSENT_HEADINGS = {
    "Key Information",
    "Why am I being invited to take part in a research study?",
    "Why is this research being done?",
    "How long will the research last and what will I need to do?",
    "Is there any way being in this study could be bad for me?",
    "Will being in this study help me in any way?",
    "What happens if I do not want to be in this research?",
    "What should I know about a research study?",
    "Who can I talk to?",
    "How many people will be studied?",
    "What happens if I say yes, I want to be in this research?",
    "What are my responsibilities if I take part in this research?",
    "What happens if I say yes, but I change my mind later?",
    "Is there any way being in this study could be bad for me? (Detailed Risks)",
    "What happens to the information collected for the research?",
    "Can I be removed from the research without my OK?",
    "What else do I need to know?",
    "Voice-Cloning Authorization",
    "Consent and Signature",
}

SUPPORT_HEADINGS = {
    "Recruitment_Materials_DRAFT.docx": {
        "Recruitment Email", "Flyer or Research-Portal Listing", "Verbal Referral Script", "Recruitment Controls"
    },
    "Eligibility_and_Safety_Screening_DRAFT.docx": {
        "Basic Eligibility", "Headset and Visual Safety", "Voice Recording Safety and Privacy Confirmation", "Pre-Session Symptom Rating", "Researcher Determination"
    },
    "Participant_Questionnaires_DRAFT.docx": {
        "Background and Prior Experience (once per participant)", "Raw NASA Task Load Index (after each block)", "Post-Block Assistance and Voice Check", "Post-Session Comparison", "Post-Session Symptom Rating", "Instrument Decisions Before Submission"
    },
    "Voice_Recording_Script_DRAFT.docx": {
        "Researcher Setup", "Participant Introduction", "Recording Passage", "Quality and Upload Record"
    },
    "Researcher_Session_and_Event_Record_DRAFT.docx": {
        "Pre-Session Controls", "Consent and Eligibility", "Voice Procedure", "Headset, Calibration, and Practice", "Experimental Blocks", "Data Transfer and Cleanup", "Session Disposition and Data-Quality Reason Codes", "Symptoms or Adverse Event", "Protocol Deviation or Confidentiality Incident"
    },
    "Post_Participation_Information_DRAFT.docx": {"Voice and Data Cleanup", "After the Headset", "Contacts"},
}

FORBIDDEN_RESIDUE = (
    "Motor Control",
    "Timeflow",
    "Meta Quest",
    "$15",
    "300 participants",
)

REVISION_ELEMENTS = {
    "ins", "del", "moveFrom", "moveTo", "moveFromRangeStart", "moveFromRangeEnd",
    "moveToRangeStart", "moveToRangeEnd", "customXmlInsRangeStart", "customXmlInsRangeEnd",
    "customXmlDelRangeStart", "customXmlDelRangeEnd", "customXmlMoveFromRangeStart",
    "customXmlMoveFromRangeEnd", "customXmlMoveToRangeStart", "customXmlMoveToRangeEnd",
    "rPrChange", "pPrChange", "sectPrChange", "tblPrChange", "tblGridChange", "trPrChange",
    "tcPrChange", "numberingChange", "cellIns", "cellDel", "cellMerge",
}
WORDPROCESSINGML_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_2010_NAMESPACE = "http://schemas.microsoft.com/office/word/2010/wordml"
REVISION_QNAMES = {
    *(f"{{{WORDPROCESSINGML_NAMESPACE}}}{name}" for name in REVISION_ELEMENTS),
    f"{{{WORD_2010_NAMESPACE}}}conflictIns",
    f"{{{WORD_2010_NAMESPACE}}}conflictDel",
}


def document_text(path: Path) -> str:
    document = Document(path)
    blocks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for part in (section.header, section.footer):
            blocks.extend(p.text for p in part.paragraphs)
            for table in part.tables:
                for row in table.rows:
                    blocks.extend(cell.text for cell in row.cells)
    return "\n".join(blocks)


def headings(path: Path) -> set[str]:
    document = Document(path)
    return {
        re.sub(r"^\d+\.\d+\s+", "", p.text.strip())
        for p in document.paragraphs
        if p.style and p.style.name.startswith("Heading")
    }


def open_decisions(text: str) -> Counter[str]:
    return Counter(
        " ".join(match.group(1).split()).strip("* ")
        for match in re.finditer(r"(?:\[?OPEN DECISION:?\s*)(.*?)(?:\]|\*\*|$)", text, re.MULTILINE)
        if match.group(1).strip("* ")
    )


def contains_revision(xml: bytes) -> bool:
    try:
        root = ET.fromstring(xml)
    except ET.ParseError:
        return True
    return any(element.tag in REVISION_QNAMES for element in root.iter())


def verify_package(path: Path) -> list[str]:
    failures: list[str] = []
    text = document_text(path)
    if "DRAFT — NOT FOR SUBMISSION OR PARTICIPANT USE" not in text:
        failures.append(f"{path.name}: missing draft-use warning")
    source_decisions = open_decisions((SOURCE / SOURCE_FOR_OUTPUT[path.name]).read_text(encoding="utf-8"))
    output_decisions = open_decisions(text)
    for decision, expected_count in source_decisions.items():
        if output_decisions[decision] < expected_count:
            failures.append(f"{path.name}: generated output lost OPEN DECISION text: {decision}")
    for residue in FORBIDDEN_RESIDUE:
        if residue.casefold() in text.casefold():
            failures.append(f"{path.name}: contains residue from prior study: {residue}")

    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
        if any(name.startswith("word/comments") for name in names):
            failures.append(f"{path.name}: contains Word comments")
        for name in names:
            if not name.endswith(".xml"):
                continue
            xml = package.read(name)
            if contains_revision(xml):
                failures.append(f"{path.name}: contains tracked changes in {name}")
                break
    return failures


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    output = args.output_dir.resolve()
    failures: list[str] = []
    actual = {path.name for path in output.glob("*.docx")}
    missing = EXPECTED - actual
    unexpected = actual - EXPECTED
    if missing:
        failures.append(f"Missing outputs: {', '.join(sorted(missing))}")
    if unexpected:
        failures.append(f"Unexpected outputs: {', '.join(sorted(unexpected))}")

    for filename in sorted(EXPECTED & actual):
        failures.extend(verify_package(output / filename))

    protocol = output / "HRP-503_Gaze_Contingency_Protocol_DRAFT.docx"
    consent = output / "HRP-502_Gaze_Contingency_Adult_Consent_DRAFT.docx"
    if protocol.exists():
        missing_headings = PROTOCOL_HEADINGS - headings(protocol)
        if missing_headings:
            failures.append(f"Protocol missing headings: {', '.join(sorted(missing_headings))}")
    if consent.exists():
        missing_headings = CONSENT_HEADINGS - headings(consent)
        if missing_headings:
            failures.append(f"Consent missing headings: {', '.join(sorted(missing_headings))}")
    for filename, required in SUPPORT_HEADINGS.items():
        path = output / filename
        if path.exists():
            missing_headings = required - headings(path)
            if missing_headings:
                failures.append(f"{filename} missing headings: {', '.join(sorted(missing_headings))}")

    if failures:
        raise SystemExit("IRB packet verification failed:\n- " + "\n- ".join(failures))
    print(f"Verified {len(EXPECTED)} generated IRB draft documents.")


if __name__ == "__main__":
    main()
